"""
Document processing engine for legal document generation
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document

try:
    from .config import (
        TEMPLATES_DIR, OUTPUT_DIR, COMPANY_INFO, DEFAULT_AGREEMENT,
        NAMING_CONVENTIONS, VALIDATION_PATTERNS
    )
except ImportError:
    from config import (
        TEMPLATES_DIR, OUTPUT_DIR, COMPANY_INFO, DEFAULT_AGREEMENT,
        NAMING_CONVENTIONS, VALIDATION_PATTERNS
    )


class DocumentProcessor:
    """Processes and generates legal documents from templates"""
    
    def __init__(self):
        self.templates_dir = Path(TEMPLATES_DIR)
        self.output_dir = Path(OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_nda(self, individual_info: Dict, project_type: str, 
                     agreement_date: Optional[str] = None) -> Tuple[str, Dict]:
        """Generate NDA document for individual"""
        
        # Validate inputs
        self._validate_individual_info(individual_info)
        
        # Prepare document data
        doc_data = self._prepare_document_data(individual_info, project_type, agreement_date)
        
        # Load template
        template_path = self.templates_dir / "NDA Template.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"NDA template not found: {template_path}")
        
        # Generate output filename
        output_filename = self._generate_filename("nda", individual_info)
        output_path = self.output_dir / output_filename
        
        # Process document
        start_time = datetime.now()
        self._process_document(template_path, output_path, doc_data)
        end_time = datetime.now()
        
        # Verify completion
        verification_result = self._verify_document(output_path)
        
        # Prepare summary
        summary = {
            "input_template": "NDA Template.docx",
            "output_file": str(output_path),
            "individual": individual_info.get("full_name"),
            "email": individual_info.get("email"),
            "project_type": project_type,
            "tool_used": "python-docx + template preservation",
            "start_time": start_time.strftime("%Y-%m-%d %H:%M:%S"),
            "completion_time": end_time.strftime("%Y-%m-%d %H:%M:%S"),
            "duration_seconds": (end_time - start_time).total_seconds(),
            "status": "✅ Success" if verification_result["success"] else "❌ Issues found",
            "placeholders_remaining": verification_result["placeholders_found"],
            "file_size_kb": round(output_path.stat().st_size / 1024, 2) if output_path.exists() else 0
        }
        
        return str(output_path), summary
    
    def _validate_individual_info(self, individual_info: Dict) -> None:
        """Validate individual information completeness"""
        
        required_fields = [
            "full_name", "address", "email", "professional_description"
        ]
        
        missing_fields = [field for field in required_fields if not individual_info.get(field)]
        
        if missing_fields:
            raise ValueError(f"Missing required individual information: {', '.join(missing_fields)}")
        
        # Validate email format
        email_pattern = VALIDATION_PATTERNS["email"]
        if not re.match(email_pattern, individual_info["email"]):
            raise ValueError(f"Invalid email format: {individual_info['email']}")
    
    def _prepare_document_data(self, individual_info: Dict, project_type: str, 
                             agreement_date: Optional[str]) -> Dict:
        """Prepare all data needed for document generation"""
        
        # Parse agreement date
        if agreement_date:
            try:
                parsed_date = datetime.strptime(agreement_date, "%Y-%m-%d")
                date_day = parsed_date.strftime("%d").lstrip("0") + self._get_ordinal_suffix(parsed_date.day)
                date_month = parsed_date.strftime("%B")
                date_year = parsed_date.strftime("%Y")
                formatted_date = f"{date_day} day of {date_month}, {date_year}"
                effective_date = f"{date_day} {date_month}, {date_year}"
            except ValueError:
                # Default to today if parsing fails
                today = datetime.now()
                date_day = today.strftime("%d").lstrip("0") + self._get_ordinal_suffix(today.day)
                date_month = today.strftime("%B")
                date_year = today.strftime("%Y")
                formatted_date = f"{date_day} day of {date_month}, {date_year}"
                effective_date = f"{date_day} {date_month}, {date_year}"
        else:
            # Default to today
            today = datetime.now()
            date_day = today.strftime("%d").lstrip("0") + self._get_ordinal_suffix(today.day)
            date_month = today.strftime("%B") 
            date_year = today.strftime("%Y")
            formatted_date = f"{date_day} day of {date_month}, {date_year}"
            effective_date = f"{date_day} {date_month}, {date_year}"
        
        # Prepare individual short name
        individual_short_name = individual_info.get("short_name") or individual_info["full_name"].split()[0]
        
        return {
            "company": COMPANY_INFO,
            "individual": {
                **individual_info,
                "short_name": individual_short_name
            },
            "agreement": {
                **DEFAULT_AGREEMENT,
                "formatted_date": formatted_date,
                "effective_date": effective_date,
                "project_description": project_type
            }
        }
    
    def _get_ordinal_suffix(self, day: int) -> str:
        """Get ordinal suffix for day (1st, 2nd, 3rd, 4th, etc.)"""
        if 10 <= day % 100 <= 20:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        return suffix
    
    def _generate_filename(self, doc_type: str, individual_info: Dict) -> str:
        """Generate output filename based on naming conventions"""
        
        template = NAMING_CONVENTIONS.get(doc_type, "{prefix}{first_name}_{last_name}.docx")
        
        # Extract first and last name
        name_parts = individual_info["full_name"].strip().split()
        first_name = name_parts[0] if name_parts else "Unknown"
        last_name = name_parts[-1] if len(name_parts) > 1 else ""
        
        # Clean names for filename
        first_name = re.sub(r'[^\w\-_]', '', first_name)
        last_name = re.sub(r'[^\w\-_]', '', last_name)
        
        filename = template.format(
            prefix="NDA_",
            first_name=first_name,
            last_name=last_name
        )
        
        return filename
    
    def _process_document(self, template_path: Path, output_path: Path, doc_data: Dict) -> None:
        """Process template and generate final document"""
        
        # Load template
        doc = Document(template_path)
        
        # Get all paragraphs including from tables
        all_paragraphs = []
        
        for paragraph in doc.paragraphs:
            all_paragraphs.append(paragraph)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_paragraphs.append(paragraph)
        
        # Process all paragraphs with systematic replacement
        for paragraph in all_paragraphs:
            self._replace_paragraph_content(paragraph, doc_data)
        
        # Save processed document
        doc.save(output_path)
    
    def _replace_paragraph_content(self, paragraph, doc_data: Dict) -> None:
        """Replace placeholder content in paragraph while preserving formatting"""
        
        original_text = paragraph.text
        new_text = original_text
        
        company = doc_data["company"]
        individual = doc_data["individual"] 
        agreement = doc_data["agreement"]
        
        # Comprehensive single-pass replacement
        replacements = [
            # Date and venue replacements
            (f'this [●] day of October, 2025', f'this {agreement["formatted_date"]}'),
            ('which shall mean the date [●]', f'which shall mean the date {agreement["effective_date"]}'),
            ('The seat and venue of arbitration shall be [●]', f'The seat and venue of arbitration shall be {agreement["arbitration_venue"]}'),
            ('courts situated at [●]', f'courts situated at {agreement["jurisdiction"]}'),
            ('Project means [●]', f'Project means {agreement["project_description"]}'),
            
            # Descriptive paragraphs
            ('[Insert an introductory paragraph about the Disclosing Party]', company["description"]),
            ('[Insert an introductory paragraph about the Receiving Party]', individual["professional_description"]),
            
            # Party collective references
            ('[●]/Disclosing Party and [●]/Receiving Party', f'{company["short_name"]}/Disclosing Party and {individual["short_name"]}/Receiving Party'),
            
            # Confidential information sharing
            ('shared by [●] with [●] during', f'shared by {company["short_name"]} with {individual["short_name"]} during'),
        ]
        
        # Apply basic replacements
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        
        # Handle complex party definitions with context
        if '[●], having its head office at [●]' in new_text:
            if 'Disclosing Party' in new_text:
                new_text = new_text.replace('[●], having its head office at [●]', f'{company["name"]}, having its head office at {company["address"]}')
                new_text = new_text.replace('"[●]" or "Disclosing Party"', f'"{company["short_name"]}" or "Disclosing Party"')
            elif 'Receiving Party' in new_text:
                new_text = new_text.replace('[●], having its head office at [●]', f'{individual["full_name"]}, having his address at {individual["address"]}')
                new_text = new_text.replace('"[●]" or "Receiving Party"', f'"{individual["short_name"]}" or "Receiving Party"')
        
        # Handle contact information with comprehensive context detection
        context_indicators = {
            'company': ['If to the Disclosing Party', 'Disclosing Party', 'Mirai360', 'shivang@mirai360.ai'],
            'individual': ['If to the Receiving Party', 'Receiving Party']
        }
        
        contact_context = 'individual'  # default
        for context_type, indicators in context_indicators.items():
            if any(indicator in new_text for indicator in indicators):
                contact_context = context_type
                break
        
        # Apply contact information based on context
        contact_replacements = {
            'company': [
                ('Address\t\t: \t[●]', f'Address\t\t: \t{company["address"]}'),
                ('Kind Attention\t: \t[●]', f'Kind Attention\t: \t{company["kind_attention"]}'),
                ('Email\t\t\t: \t[●]', f'Email\t\t\t: \t{company["email"]}'),
            ],
            'individual': [
                ('Address\t\t: \t[●]', f'Address\t\t: \t{individual["address"]}'),
                ('Kind Attention\t: \t[●]', f'Kind Attention\t: \t{individual["full_name"]}'),
                ('Email\t\t\t: \t[●]', f'Email\t\t\t: \t{individual["email"]}'),
            ]
        }
        
        for old, new in contact_replacements[contact_context]:
            new_text = new_text.replace(old, new)
        
        # Handle signature blocks
        if 'For and on behalf of [●]' in new_text:
            if 'Disclosing' in new_text or 'Mirai360' in new_text:
                new_text = new_text.replace('For and on behalf of [●]', f'For and on behalf of {company["short_name"]}')
            else:
                new_text = new_text.replace('For and on behalf of [●]', f'For and on behalf of {individual["full_name"]}')
        
        # Final cleanup - handle any remaining [●] placeholders
        remaining_placeholders = {
            '[●]': {
                'company_contexts': ['Disclosing', 'Mirai360', 'shivang@mirai360.ai'],
                'company_value': company["short_name"],
                'individual_value': individual["short_name"]
            }
        }
        
        for placeholder, config in remaining_placeholders.items():
            if placeholder in new_text:
                if any(ctx in new_text for ctx in config['company_contexts']):
                    new_text = new_text.replace(placeholder, config['company_value'])
                else:
                    new_text = new_text.replace(placeholder, config['individual_value'])
        
        # Update paragraph text if changed
        if new_text != original_text:
            paragraph.text = new_text
    
    def _get_context_clues(self, paragraph, doc_data: Dict) -> str:
        """Get context clues to determine which entity information to use"""
        
        # Simple heuristic: if company name or Disclosing Party mentioned nearby, use company info
        text_context = paragraph.text.lower()
        
        if any(term in text_context for term in ['disclosing', 'mirai360', 'shivang']):
            return 'company'
        else:
            return 'individual'
    
    def _verify_document(self, output_path: Path) -> Dict:
        """Verify document completion by checking for remaining placeholders"""
        
        if not output_path.exists():
            return {"success": False, "error": "Output file not created", "placeholders_found": []}
        
        try:
            doc = Document(output_path)
            placeholders_found = []
            
            # Check all paragraphs
            all_paragraphs = []
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            all_paragraphs.append(paragraph)
            
            # Look for remaining placeholders
            for paragraph in all_paragraphs:
                text = paragraph.text
                for pattern in VALIDATION_PATTERNS["placeholders"]:
                    if re.search(pattern, text):
                        placeholders_found.append(text[:100] + "..." if len(text) > 100 else text)
            
            return {
                "success": len(placeholders_found) == 0,
                "placeholders_found": placeholders_found,
                "total_paragraphs": len(all_paragraphs)
            }
            
        except Exception as e:
            return {"success": False, "error": str(e), "placeholders_found": []}