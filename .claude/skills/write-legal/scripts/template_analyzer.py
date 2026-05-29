"""
Template analysis and change detection for legal documents
"""

import os
import hashlib
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Set, Tuple, Optional
from docx import Document

try:
    from .config import TEMPLATES_DIR, SKILL_DATA_DIR, VALIDATION_PATTERNS
except ImportError:
    from config import TEMPLATES_DIR, SKILL_DATA_DIR, VALIDATION_PATTERNS


class TemplateAnalyzer:
    """Analyzes legal document templates and detects changes"""
    
    def __init__(self):
        self.data_dir = Path(SKILL_DATA_DIR)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.cache_file = self.data_dir / "template_cache.json"
        self._load_cache()
    
    def _load_cache(self) -> None:
        """Load template analysis cache"""
        try:
            if self.cache_file.exists():
                with open(self.cache_file, 'r') as f:
                    self.cache = json.load(f)
            else:
                self.cache = {}
        except Exception:
            self.cache = {}
    
    def _save_cache(self) -> None:
        """Save template analysis cache"""
        try:
            with open(self.cache_file, 'w') as f:
                json.dump(self.cache, f, indent=2)
        except Exception as e:
            print(f"Warning: Could not save template cache: {e}")
    
    def get_template_signature(self, template_path: Path) -> str:
        """Generate unique signature for template file"""
        if not template_path.exists():
            return ""
        
        stat = template_path.stat()
        return f"{stat.st_mtime}_{stat.st_size}"
    
    def has_template_changed(self, template_name: str, template_path: Path) -> bool:
        """Check if template has changed since last analysis"""
        current_signature = self.get_template_signature(template_path)
        cached_signature = self.cache.get(template_name, {}).get("signature", "")
        
        return current_signature != cached_signature
    
    def analyze_template(self, template_name: str, template_path: Path) -> Dict:
        """Analyze template and extract all placeholders and structure"""
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Check if analysis is cached and template hasn't changed
        if not self.has_template_changed(template_name, template_path):
            cached_analysis = self.cache.get(template_name, {}).get("analysis")
            if cached_analysis:
                return cached_analysis
        
        # Silently analyze template
        
        try:
            doc = Document(template_path)
            analysis = self._extract_template_info(doc)
            
            # Cache the analysis
            self.cache[template_name] = {
                "signature": self.get_template_signature(template_path),
                "analysis": analysis,
                "last_analyzed": datetime.now().isoformat()
            }
            self._save_cache()
            
            return analysis
            
        except Exception as e:
            raise Exception(f"Failed to analyze template {template_name}: {e}")
    
    def _extract_template_info(self, doc: Document) -> Dict:
        """Extract placeholders and categorize them"""
        
        placeholders = set()
        all_text = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
                if self._contains_placeholder(text):
                    placeholders.add(text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            all_text.append(text)
                            if self._contains_placeholder(text):
                                placeholders.add(text)
        
        # Categorize placeholders
        categories = self._categorize_placeholders(placeholders)
        
        return {
            "total_placeholders": len(placeholders),
            "placeholders": list(placeholders),
            "categories": categories,
            "analysis_date": datetime.now().isoformat(),
            "document_structure": {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables)
            }
        }
    
    def _contains_placeholder(self, text: str) -> bool:
        """Check if text contains any placeholder patterns"""
        import re
        for pattern in VALIDATION_PATTERNS["placeholders"]:
            if re.search(pattern, text):
                return True
        return False
    
    def _categorize_placeholders(self, placeholders: Set[str]) -> Dict[str, List[str]]:
        """Categorize placeholders by type"""
        
        categories = {
            "company_info": [],
            "individual_info": [],
            "agreement_details": [],
            "contact_info": [],
            "project_info": [],
            "signature_blocks": [],
            "other": []
        }
        
        for placeholder in placeholders:
            lower_text = placeholder.lower()
            
            if any(term in lower_text for term in ["disclosing party", "head office"]) and "receiving party" not in lower_text:
                categories["company_info"].append(placeholder)
            elif "receiving party" in lower_text:
                categories["individual_info"].append(placeholder)
            elif any(term in lower_text for term in ["date", "arbitration", "court", "jurisdiction"]):
                categories["agreement_details"].append(placeholder)
            elif any(term in lower_text for term in ["address", "email", "kind attention"]):
                categories["contact_info"].append(placeholder)
            elif any(term in lower_text for term in ["project means", "purpose"]):
                categories["project_info"].append(placeholder)
            elif "for and on behalf" in lower_text:
                categories["signature_blocks"].append(placeholder)
            else:
                categories["other"].append(placeholder)
        
        return categories
    
    def get_analysis_summary(self, template_name: str) -> Optional[Dict]:
        """Get cached analysis summary for template"""
        return self.cache.get(template_name, {}).get("analysis")
    
    def get_required_fields(self, template_name: str) -> List[str]:
        """Get list of required information fields for template"""
        analysis = self.get_analysis_summary(template_name)
        if not analysis:
            return []
        
        required_fields = []
        
        # Based on placeholder categories, determine required fields
        categories = analysis.get("categories", {})
        
        if categories.get("individual_info"):
            required_fields.extend([
                "full_name", "address", "email", "professional_description", "short_name"
            ])
        
        if categories.get("project_info"):
            required_fields.append("project_type")
        
        if categories.get("agreement_details"):
            required_fields.extend([
                "agreement_date", "jurisdiction", "arbitration_venue"
            ])
        
        return required_fields