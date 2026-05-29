"""
Main write_legal skill implementation
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Add current directory to path for imports
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))

# Handle both relative and absolute imports
try:
    from .config import (
        TEMPLATES_DIR, OUTPUT_DIR, PROJECT_TYPES, SUPPORTED_TEMPLATES, 
        COMPANY_INFO, PERFORMANCE_TARGETS
    )
    from .template_analyzer import TemplateAnalyzer
    from .individual_database import IndividualDatabase
    from .document_processor import DocumentProcessor
except ImportError:
    # Fallback to absolute imports
    from config import (
        TEMPLATES_DIR, OUTPUT_DIR, PROJECT_TYPES, SUPPORTED_TEMPLATES, 
        COMPANY_INFO, PERFORMANCE_TARGETS
    )
    from template_analyzer import TemplateAnalyzer
    from individual_database import IndividualDatabase
    from document_processor import DocumentProcessor


class WriteLegalSkill:
    """Main write_legal skill class for legal document generation"""
    
    def __init__(self):
        self.template_analyzer = TemplateAnalyzer()
        self.individual_db = IndividualDatabase()
        self.doc_processor = DocumentProcessor()
        self._ensure_environment()
    
    def _ensure_environment(self) -> None:
        """Ensure required environment and dependencies"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )
        
        # Ensure directories exist
        Path(TEMPLATES_DIR).mkdir(parents=True, exist_ok=True)
        Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    
    def create_nda(self, input_text: str = None, **kwargs) -> Dict:
        """
        Main entry point for NDA creation
        
        Args:
            input_text: Freeform text with individual information
            **kwargs: Direct field specification (full_name, email, etc.)
        
        Returns:
            Dict with creation results and summary
        """
        
        start_time = datetime.now()
        
        try:
            # Step 1: Check template status
            template_info = self._check_template_status("nda")
            
            # Step 2: Gather individual information
            individual_info = self._gather_individual_info(input_text, **kwargs)
            
            # Step 3: Get project type
            project_type = self._get_project_type(kwargs.get("project_type"))
            
            # Step 4: Get agreement details
            agreement_date = kwargs.get("agreement_date")
            
            # Step 5: Generate document
            output_path, summary = self.doc_processor.generate_nda(
                individual_info, project_type, agreement_date
            )
            
            # Step 6: Update database
            self.individual_db.add_individual(individual_info)
            self.individual_db.update_document_count(individual_info["email"], project_type)
            
            # Step 7: Prepare final summary
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            final_summary = {
                **summary,
                "template_status": template_info,
                "performance": {
                    "total_duration": duration,
                    "target_duration": PERFORMANCE_TARGETS["total_time"],
                    "within_target": duration <= PERFORMANCE_TARGETS["total_time"]
                }
            }
            
            return {
                "success": True,
                "output_file": output_path,
                "summary": final_summary,
                "individual_profile": individual_info
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "duration": (datetime.now() - start_time).total_seconds()
            }
    
    def _check_template_status(self, doc_type: str) -> Dict:
        """Check template status and analyze if needed"""
        
        if doc_type not in SUPPORTED_TEMPLATES:
            raise ValueError(f"Unsupported document type: {doc_type}")
        
        template_config = SUPPORTED_TEMPLATES[doc_type]
        
        if not template_config["supported"]:
            raise ValueError(
                f"{doc_type.upper()} documents are not yet supported. "
                f"Currently supported: {[k for k, v in SUPPORTED_TEMPLATES.items() if v['supported']]}"
            )
        
        template_path = Path(TEMPLATES_DIR) / template_config["template_file"]
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Analyze template if needed
        analysis = self.template_analyzer.analyze_template(doc_type, template_path)
        
        return {
            "template_file": template_config["template_file"],
            "template_path": str(template_path),
            "analysis_date": analysis.get("analysis_date"),
            "total_placeholders": analysis.get("total_placeholders"),
            "categories": list(analysis.get("categories", {}).keys())
        }
    
    def _gather_individual_info(self, input_text: str = None, **kwargs) -> Dict:
        """Gather and validate individual information"""
        
        individual_info = {}
        
        # Start with kwargs (direct field specification)
        individual_info.update(kwargs)
        
        # Parse freeform input if provided
        if input_text:
            parsed_info = self.individual_db.parse_freeform_input(input_text)
            
            # Merge parsed info (kwargs take precedence)
            for key, value in parsed_info.items():
                if key not in individual_info:
                    individual_info[key] = value
        
        # Try to get existing profile for auto-completion
        if individual_info.get("email"):
            existing_profile = self.individual_db.get_individual(individual_info["email"])
            if existing_profile:
                # Merge existing profile (new info takes precedence)
                merged_info = {**existing_profile, **individual_info}
                individual_info = merged_info
        
        # Validate and prompt for missing required fields
        required_fields = ["full_name", "address", "email", "professional_description"]
        missing_fields = [field for field in required_fields if not individual_info.get(field)]
        
        if missing_fields:
            # Show checklist of missing information
            checklist = self._create_information_checklist(individual_info, missing_fields)
            raise ValueError(f"Missing required information:\\n{checklist}")
        
        return individual_info
    
    def _create_information_checklist(self, current_info: Dict, missing_fields: List[str]) -> str:
        """Create a checklist showing current and missing information"""
        
        checklist_items = [
            ("full_name", "✅ Full Legal Name" if current_info.get("full_name") else "❌ Full Legal Name"),
            ("address", "✅ Complete Address" if current_info.get("address") else "❌ Complete Address"),
            ("email", "✅ Email Address" if current_info.get("email") else "❌ Email Address"),
            ("professional_description", "✅ Professional Description" if current_info.get("professional_description") else "❌ Professional Description"),
            ("short_name", "✅ Short Reference Name" if current_info.get("short_name") else "⚠️ Short Reference Name (optional)"),
        ]
        
        checklist = "\\n📋 INFORMATION CHECKLIST:\\n"
        for field, status in checklist_items:
            checklist += f"   {status}\\n"
            if current_info.get(field):
                checklist += f"      → {current_info[field]}\\n"
        
        if missing_fields:
            checklist += f"\\n❌ Please provide: {', '.join(missing_fields)}"
        
        # Show suggestions if we have partial information
        if current_info.get("email") or current_info.get("full_name"):
            query = current_info.get("email") or current_info.get("full_name")
            suggestions = self.individual_db.get_suggestions(query)
            
            if suggestions:
                checklist += "\\n\\n💡 FOUND EXISTING PROFILES:"
                for suggestion in suggestions[:3]:
                    checklist += f"\\n   • {suggestion['full_name']} ({suggestion['email']}) - {suggestion['documents_count']} docs"
        
        return checklist
    
    def _get_project_type(self, specified_type: str = None) -> str:
        """Get project type with selection or custom input"""
        
        if specified_type:
            # Validate against known types or accept custom
            if specified_type in PROJECT_TYPES:
                return specified_type
            elif specified_type == "Custom - User Specified":
                raise ValueError("Please specify the custom project type")
            else:
                # Assume it's a custom type
                return specified_type
        
        # If not specified, show available options
        options_text = "\\n".join([f"{i+1}. {ptype}" for i, ptype in enumerate(PROJECT_TYPES)])
        
        raise ValueError(
            f"Please specify project type:\\n\\n{options_text}\\n\\n"
            f"Use project_type parameter with one of the above options or specify your custom type."
        )
    
    def get_individual_suggestions(self, query: str) -> List[Dict]:
        """Get individual profile suggestions for auto-completion"""
        return self.individual_db.get_suggestions(query)
    
    def get_project_types(self) -> List[str]:
        """Get available project types"""
        return PROJECT_TYPES.copy()
    
    def get_database_stats(self) -> Dict:
        """Get database statistics"""
        return self.individual_db.get_database_stats()
    
    def validate_template(self, doc_type: str = "nda") -> Dict:
        """Validate template and return analysis"""
        
        try:
            template_info = self._check_template_status(doc_type)
            return {
                "success": True,
                "template_info": template_info
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def edit_existing_docx(self, file_path: str, edits: List[Dict], output_path: str = None) -> Dict:
        """Edit existing docx file directly without format conversion
        
        Args:
            file_path: Path to existing docx file
            edits: List of edit operations [{"find": "text", "replace": "new_text"}]
            output_path: Optional output path (defaults to same as input)
            
        Returns:
            Dict with edit results
        """
        
        start_time = datetime.now()
        
        try:
            from docx import Document
            import shutil
            
            input_path = Path(file_path)
            if not input_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine output path
            if output_path:
                final_output = Path(output_path)
            else:
                final_output = input_path
            
            # Create backup copy first
            backup_path = input_path.parent / f"{input_path.stem}_backup{input_path.suffix}"
            shutil.copy2(input_path, backup_path)
            
            # Load document
            doc = Document(input_path)
            
            # Track changes
            changes_made = []
            
            # Apply edits to paragraphs
            for edit in edits:
                find_text = edit.get("find", "")
                replace_text = edit.get("replace", "")
                
                if not find_text:
                    continue
                
                # Search and replace in paragraphs
                for paragraph in doc.paragraphs:
                    if find_text in paragraph.text:
                        original_text = paragraph.text
                        new_text = paragraph.text.replace(find_text, replace_text)
                        
                        # Clear paragraph and add new text
                        paragraph.clear()
                        paragraph.add_run(new_text)
                        
                        changes_made.append({
                            "location": "paragraph",
                            "original": original_text[:100] + "..." if len(original_text) > 100 else original_text,
                            "modified": new_text[:100] + "..." if len(new_text) > 100 else new_text
                        })
                
                # Search and replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if find_text in cell.text:
                                original_text = cell.text
                                new_text = cell.text.replace(find_text, replace_text)
                                
                                # Clear cell and add new text
                                cell._tc.clear_content()
                                cell.text = new_text
                                
                                changes_made.append({
                                    "location": "table_cell",
                                    "original": original_text[:100] + "..." if len(original_text) > 100 else original_text,
                                    "modified": new_text[:100] + "..." if len(new_text) > 100 else new_text
                                })
            
            # Save document
            doc.save(final_output)
            
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            return {
                "success": True,
                "input_file": str(input_path),
                "output_file": str(final_output),
                "backup_file": str(backup_path),
                "changes_made": changes_made,
                "total_changes": len(changes_made),
                "duration_seconds": duration,
                "completion_time": end_time.strftime("%Y-%m-%d %H:%M:%S")
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "duration_seconds": (datetime.now() - start_time).total_seconds()
            }
    
    def copy_and_rename_docx(self, source_path: str, new_name: str, output_dir: str = None) -> Dict:
        """Copy and rename docx file for editing
        
        Args:
            source_path: Path to source docx file
            new_name: New filename (without extension)
            output_dir: Optional output directory (defaults to source directory)
            
        Returns:
            Dict with copy results
        """
        
        try:
            import shutil
            
            source = Path(source_path)
            if not source.exists():
                raise FileNotFoundError(f"Source file not found: {source_path}")
            
            # Determine output directory
            if output_dir:
                target_dir = Path(output_dir)
            else:
                target_dir = source.parent
            
            # Ensure target directory exists
            target_dir.mkdir(parents=True, exist_ok=True)
            
            # Create target path
            target_path = target_dir / f"{new_name}.docx"
            
            # Copy file
            shutil.copy2(source, target_path)
            
            return {
                "success": True,
                "source_file": str(source),
                "target_file": str(target_path),
                "operation": "copy_and_rename"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def format_completion_summary(self, summary: Dict, verbose: bool = False) -> str:
        """Format completion summary for display"""
        
        if not summary:
            return "❌ No summary available"
        
        # Simple format by default
        if not verbose:
            status_icon = "✅" if "Success" in summary.get('status', '') else "❌"
            return f"{status_icon} NDA created: {summary.get('output_file', 'Unknown')}"
        
        # Detailed format when requested
        duration_min = round(summary.get("duration_seconds", 0) / 60, 1)
        performance_icon = "🚀" if summary.get("performance", {}).get("within_target", False) else "⏱️"
        
        formatted_summary = f"""
📋 EXECUTION SUMMARY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Output File: {summary.get('output_file', 'Unknown')}
Individual: {summary.get('individual', 'Unknown')} ({summary.get('email', 'Unknown')})
Project: {summary.get('project_type', 'Unknown')}
Completed: {summary.get('completion_time', 'Unknown')} ({duration_min} minutes) {performance_icon}
Status: {summary.get('status', 'Unknown')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        
        if summary.get('placeholders_remaining'):
            formatted_summary += f"⚠️ {len(summary['placeholders_remaining'])} placeholders may need attention\\n"
        
        return formatted_summary.strip()


def create_nda(input_text: str = None, **kwargs) -> Dict:
    """
    Convenience function for creating NDAs
    
    Usage examples:
        create_nda("NDA for John Doe, john@example.com, backend developer")
        create_nda(full_name="John Doe", email="john@example.com", project_type="Full Stack Development Services")
    """
    
    skill = WriteLegalSkill()
    return skill.create_nda(input_text, **kwargs)


def edit_docx(file_path: str, edits: List[Dict], output_path: str = None) -> Dict:
    """
    Convenience function for editing existing docx files
    
    Usage examples:
        edit_docx("/path/to/file.docx", [{"find": "old text", "replace": "new text"}])
        edit_docx("/path/to/file.docx", [{"find": "Recipient", "replace": "Advisor"}], "/path/to/output.docx")
    """
    
    skill = WriteLegalSkill()
    return skill.edit_existing_docx(file_path, edits, output_path)


def copy_rename_docx(source_path: str, new_name: str, output_dir: str = None) -> Dict:
    """
    Convenience function for copying and renaming docx files
    
    Usage examples:
        copy_rename_docx("/path/to/original.docx", "new_filename")
        copy_rename_docx("/path/to/original.docx", "new_filename", "/different/directory")
    """
    
    skill = WriteLegalSkill()
    return skill.copy_and_rename_docx(source_path, new_name, output_dir)


# CLI interface for testing
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Write Legal - Legal Document Generation")
    parser.add_argument("--input", "-i", help="Freeform input text")
    parser.add_argument("--name", help="Full name")
    parser.add_argument("--email", help="Email address")
    parser.add_argument("--address", help="Full address")
    parser.add_argument("--description", help="Professional description")
    parser.add_argument("--project", help="Project type")
    parser.add_argument("--date", help="Agreement date (YYYY-MM-DD)")
    parser.add_argument("--validate", action="store_true", help="Validate template only")
    parser.add_argument("--stats", action="store_true", help="Show database stats")
    
    args = parser.parse_args()
    
    skill = WriteLegalSkill()
    
    if args.validate:
        result = skill.validate_template()
        print("Template Validation:", result)
    elif args.stats:
        stats = skill.get_database_stats()
        print("Database Stats:", stats)
    else:
        kwargs = {}
        if args.name:
            kwargs["full_name"] = args.name
        if args.email:
            kwargs["email"] = args.email
        if args.address:
            kwargs["address"] = args.address
        if args.description:
            kwargs["professional_description"] = args.description
        if args.project:
            kwargs["project_type"] = args.project
        if args.date:
            kwargs["agreement_date"] = args.date
        
        result = skill.create_nda(args.input, **kwargs)
        
        if result["success"]:
            print(skill.format_completion_summary(result["summary"]))
        else:
            print(f"❌ Error: {result['error']}")